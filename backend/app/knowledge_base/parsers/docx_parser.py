from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio

from .base import DocumentParser, ParsedDocument, ChunkResult


class DocxParser(DocumentParser):
    
    async def parse(self, file_path: str) -> ParsedDocument:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._parse_sync, file_path)
    
    def _parse_sync(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        
        paragraphs = []
        sections = []
        tables = []
        
        current_section = None
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if para.style.name.startswith('Heading'):
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    sections.append(current_section)
                
                level = int(para.style.name.replace('Heading ', '')) if para.style.name != 'Heading' else 1
                current_section = {
                    'level': level,
                    'title': text,
                    'content': '',
                }
                current_content = []
            else:
                paragraphs.append(text)
                current_content.append(text)
        
        if current_section:
            current_section['content'] = '\n'.join(current_content)
            sections.append(current_section)
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'data': table_data,
                'rows': len(table.rows),
                'cols': len(table.columns),
            })
        
        content = '\n\n'.join(paragraphs)
        
        metadata = {
            'filename': Path(file_path).name,
            'file_type': 'docx',
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
        }
        
        core_props = doc.core_properties
        if core_props.title:
            metadata['title'] = core_props.title
        if core_props.author:
            metadata['author'] = core_props.author
        if core_props.subject:
            metadata['subject'] = core_props.subject
        
        return ParsedDocument(
            content=content,
            metadata=metadata,
            sections=sections if sections else None,
            tables=tables if tables else None,
        )
    
    def supported_extensions(self) -> List[str]:
        return ['.docx']
    
    def chunk_with_sections(
        self,
        parsed_doc: ParsedDocument,
        chunk_size: int = 500,
        overlap: int = 50,
    ) -> List[ChunkResult]:
        chunks = []
        
        if parsed_doc.sections:
            for section in parsed_doc.sections:
                section_content = f"## {section['title']}\n\n{section['content']}"
                
                if len(section_content) <= chunk_size:
                    chunks.append(ChunkResult(
                        content=section_content,
                        token_count=self.count_tokens(section_content),
                        section_title=section['title'],
                        chunk_metadata={'level': section['level']},
                    ))
                else:
                    text_chunks = self.chunk_text(section_content, chunk_size, overlap)
                    for i, chunk in enumerate(text_chunks):
                        chunks.append(ChunkResult(
                            content=chunk,
                            token_count=self.count_tokens(chunk),
                            section_title=section['title'],
                            chunk_metadata={
                                'level': section['level'],
                                'chunk_index': i,
                            },
                        ))
        else:
            text_chunks = self.chunk_text(parsed_doc.content, chunk_size, overlap)
            for i, chunk in enumerate(text_chunks):
                chunks.append(ChunkResult(
                    content=chunk,
                    token_count=self.count_tokens(chunk),
                    chunk_metadata={'chunk_index': i},
                ))
        
        if parsed_doc.tables:
            for i, table in enumerate(parsed_doc.tables):
                table_text = self._table_to_text(table)
                if table_text:
                    chunks.append(ChunkResult(
                        content=table_text,
                        token_count=self.count_tokens(table_text),
                        chunk_metadata={'table_index': i, 'is_table': True},
                    ))
        
        return chunks
    
    def _table_to_text(self, table: Dict) -> str:
        if not table.get('data'):
            return ''
        
        lines = []
        for row in table['data']:
            lines.append(' | '.join(row))
        
        return '\n'.join(lines)
