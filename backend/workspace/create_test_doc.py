#!/usr/bin/env python3
"""
创建一个简单的测试 Word 文档
"""

# 首先检查是否安装了 python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    print("python-docx 库已安装")
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    exit(1)

def create_test_document():
    """创建测试文档"""
    
    # 创建文档对象
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('测试文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph('这是一个简单的测试文档')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加创建时间
    from datetime import datetime
    date_para = doc.add_paragraph(f'创建时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    doc.add_paragraph()
    
    # 第一部分
    doc.add_heading('第一部分：文档简介', level=1)
    doc.add_paragraph('这是文档的第一部分，主要介绍文档的基本信息和用途。')
    doc.add_paragraph('测试文档通常用于验证文档创建工具的功能是否正常。')
    
    # 添加空行
    doc.add_paragraph()
    
    # 第二部分
    doc.add_heading('第二部分：功能测试', level=1)
    doc.add_paragraph('这部分测试文档的各种功能，包括：')
    
    # 添加列表
    doc.add_paragraph('1. 标题样式', style='List Number')
    doc.add_paragraph('2. 段落文本', style='List Number')
    doc.add_paragraph('3. 表格功能', style='List Number')
    doc.add_paragraph('4. 格式设置', style='List Number')
    
    # 添加空行
    doc.add_paragraph()
    
    # 表格部分
    doc.add_heading('测试表格', level=1)
    
    # 创建表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '列1'
    header_cells[1].text = '列2'
    header_cells[2].text = '列3'
    
    # 设置表头样式
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.style = doc.styles['Heading 4']
    
    # 填充表格数据
    data = [
        ['测试数据1', '测试数据2', '测试数据3'],
        ['示例A', '示例B', '示例C'],
        ['项目1', '项目2', '项目3']
    ]
    
    for i, row_data in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加格式化文本示例
    doc.add_heading('格式示例', level=1)
    
    # 加粗文本
    bold_para = doc.add_paragraph()
    bold_run = bold_para.add_run('这是加粗文本：')
    bold_run.bold = True
    bold_para.add_run(' 这是普通文本')
    
    # 斜体文本
    italic_para = doc.add_paragraph()
    italic_run = italic_para.add_run('这是斜体文本：')
    italic_run.italic = True
    italic_para.add_run(' 这是普通文本')
    
    # 带颜色的文本
    color_para = doc.add_paragraph()
    color_run = color_para.add_run('这是红色文本：')
    color_run.font.color.rgb = RGBColor(255, 0, 0)
    color_para.add_run(' 这是普通文本')
    
    # 添加空行
    doc.add_paragraph()
    
    # 结尾
    end_para = doc.add_paragraph('--- 文档结束 ---')
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.runs[0]
    end_run.bold = True
    
    # 保存文档
    filename = 'test_document_python.docx'
    doc.save(filename)
    
    print(f"测试文档已创建：{filename}")
    print("文档包含以下内容：")
    print("1. 标题和副标题")
    print("2. 多个段落和列表")
    print("3. 带有表头的表格")
    print("4. 格式化的文本（加粗、斜体、颜色）")
    print("5. 居中对齐的文本")
    
    return filename

if __name__ == '__main__':
    create_test_document()