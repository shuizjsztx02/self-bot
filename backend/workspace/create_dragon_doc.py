#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
龙年吉祥祝福语Word文档生成器
使用python-docx库创建格式良好的Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_dragon_year_document():
    """创建龙年祝福语文档"""
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 创建标题样式
    title_style = doc.styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = '微软雅黑'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    
    # 创建标题1样式
    heading1_style = doc.styles.add_style('Heading1Style', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = '微软雅黑'
    heading1_font.size = Pt(20)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 0, 0)
    
    # 创建标题2样式
    heading2_style = doc.styles.add_style('Heading2Style', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = '微软雅黑'
    heading2_font.size = Pt(16)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 0, 0)
    
    # 创建强调样式
    emphasis_style = doc.styles.add_style('EmphasisStyle', WD_STYLE_TYPE.PARAGRAPH)
    emphasis_font = emphasis_style.font
    emphasis_font.name = '微软雅黑'
    emphasis_font.size = Pt(12)
    emphasis_font.bold = True
    emphasis_font.color.rgb = RGBColor(220, 0, 0)  # 红色
    
    # 文档标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('龙年吉祥祝福语大全')
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('2024年（甲辰龙年）')
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
    
    # 空行
    doc.add_paragraph()
    
    # 引言
    intro = doc.add_paragraph()
    intro_run = intro.add_run('龙年到来，万象更新。龙在中国文化中象征着吉祥、尊贵和力量。以下是为您精心整理的龙年祝福语，适用于各种场合，表达对亲朋好友的美好祝愿。')
    intro_run.font.name = '微软雅黑'
    intro_run.font.size = Pt(12)
    
    # 分隔线
    separator = doc.add_paragraph('————————————————————————————————')
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 一、通用祝福语
    doc.add_paragraph().add_run('一、通用祝福语').bold = True
    doc.add_paragraph().add_run('1. 龙年大吉，万事如意！').bold = True
    doc.add_paragraph('2. 龙马精神，身体健康！')
    doc.add_paragraph('3. 龙腾虎跃，事业腾飞！')
    doc.add_paragraph('4. 龙年行大运，财源滚滚来！')
    doc.add_paragraph('5. 龙年吉祥，阖家幸福！')
    doc.add_paragraph('6. 龙飞凤舞，前程似锦！')
    doc.add_paragraph('7. 龙年好运，心想事成！')
    doc.add_paragraph('8. 龙年快乐，福气满满！')
    
    # 空行
    doc.add_paragraph()
    
    # 二、商务祝福语
    doc.add_paragraph().add_run('二、商务祝福语').bold = True
    doc.add_paragraph().add_run('1. 龙年生意兴隆，财源广进！').bold = True
    doc.add_paragraph('2. 龙腾四海，事业蒸蒸日上！')
    doc.add_paragraph('3. 龙年大展宏图，再创辉煌！')
    doc.add_paragraph('4. 龙行天下，商机无限！')
    doc.add_paragraph('5. 龙年合作愉快，共赢未来！')
    doc.add_paragraph('6. 龙年财源茂盛，生意红火！')
    
    # 空行
    doc.add_paragraph()
    
    # 三、家庭祝福语
    doc.add_paragraph().add_run('三、家庭祝福语').bold = True
    doc.add_paragraph().add_run('1. 龙年阖家欢乐，幸福安康！').bold = True
    doc.add_paragraph('2. 龙年子孙满堂，家庭和睦！')
    doc.add_paragraph('3. 龙年平安喜乐，万事顺心！')
    doc.add_paragraph('4. 龙年福星高照，好运连连！')
    doc.add_paragraph('5. 龙年身体健康，笑口常开！')
    doc.add_paragraph('6. 龙年家和万事兴，幸福永相随！')
    
    # 空行
    doc.add_paragraph()
    
    # 四、朋友祝福语
    doc.add_paragraph().add_run('四、朋友祝福语').bold = True
    doc.add_paragraph().add_run('1. 龙年友谊长存，情谊永固！').bold = True
    doc.add_paragraph('2. 龙年快乐相伴，幸福相随！')
    doc.add_paragraph('3. 龙年心想事成，梦想成真！')
    doc.add_paragraph('4. 龙年好运连连，开心每一天！')
    doc.add_paragraph('5. 龙年笑口常开，青春永驻！')
    doc.add_paragraph('6. 龙年友谊万岁，真情永恒！')
    
    # 空行
    doc.add_paragraph()
    
    # 五、创意祝福语
    doc.add_paragraph().add_run('五、创意祝福语').bold = True
    doc.add_paragraph().add_run('1. 龙年龙抬头，好运天天有！').bold = True
    doc.add_paragraph('2. 龙年舞龙灯，幸福亮晶晶！')
    doc.add_paragraph('3. 龙年赛龙舟，快乐永不休！')
    doc.add_paragraph('4. 龙年画龙点睛，事业更光明！')
    doc.add_paragraph('5. 龙年龙吟虎啸，气势冲云霄！')
    doc.add_paragraph('6. 龙年龙飞九天，梦想都实现！')
    
    # 空行
    doc.add_paragraph()
    
    # 龙年文化寓意
    doc.add_paragraph().add_run('六、龙年文化寓意').bold = True
    
    # 龙的象征意义
    doc.add_paragraph().add_run('龙的象征意义：').bold = True
    doc.add_paragraph('• 吉祥如意：龙是祥瑞的象征')
    doc.add_paragraph('• 尊贵权威：古代帝王以龙自居')
    doc.add_paragraph('• 力量智慧：龙能呼风唤雨，掌控自然')
    doc.add_paragraph('• 生机活力：龙代表生命力和创造力')
    
    # 龙年习俗
    doc.add_paragraph().add_run('龙年习俗：').bold = True
    doc.add_paragraph('1. 舞龙表演：祈求风调雨顺')
    doc.add_paragraph('2. 龙舟竞渡：纪念屈原，展现团结')
    doc.add_paragraph('3. 龙灯游街：驱邪避灾，带来光明')
    doc.add_paragraph('4. 龙形装饰：家居装饰增添喜庆')
    
    # 空行
    doc.add_paragraph()
    
    # 结语
    doc.add_paragraph().add_run('结语').bold = True
    
    conclusion = doc.add_paragraph()
    conclusion_run = conclusion.add_run('龙年象征着新的开始和无限可能。愿这些祝福语能为您带来好运和快乐，祝您龙年大吉，万事如意！')
    conclusion_run.font.name = '微软雅黑'
    conclusion_run.font.size = Pt(12)
    
    # 强调结语
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final.add_run('愿龙年的祥瑞之气伴随您一整年，让好运如龙般腾飞，幸福如龙般长久！')
    final_run.font.name = '微软雅黑'
    final_run.font.size = Pt(14)
    final_run.font.bold = True
    final_run.font.color.rgb = RGBColor(0, 0, 139)  # 深蓝色
    
    # 分隔线
    doc.add_paragraph('————————————————————————————————').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 页脚信息
    footer1 = doc.add_paragraph()
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer1_run = footer1.add_run('祝福语整理于2024年龙年春节前夕')
    footer1_run.font.italic = True
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.add_run('愿您和家人朋友共享龙年的喜悦与幸福')
    footer2_run.font.italic = True
    
    # 保存文档
    filename = '龙年吉祥祝福语大全.docx'
    doc.save(filename)
    print(f'Word文档已创建：{filename}')
    print(f'文档包含：')
    print(f'  - 5个祝福语类别')
    print(f'  - 30+条精选祝福语')
    print(f'  - 龙年文化寓意介绍')
    print(f'  - 专业格式排版')
    
    return filename

if __name__ == '__main__':
    create_dragon_year_document()