#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
生成三个主题的Word文档
"""

import os
import sys
from datetime import datetime

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    print("✓ python-docx库已导入")
except ImportError as e:
    print(f"✗ 导入python-docx失败: {e}")
    print("请安装: pip install python-docx")
    sys.exit(1)

def create_document(title, content, filename):
    """创建单个Word文档"""
    try:
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(12)
        
        # 添加标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副标题
        subtitle = doc.add_paragraph('文档摘要')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加创建时间
        date_str = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        date_para = doc.add_paragraph(f'创建时间: {date_str}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分隔线
        doc.add_paragraph('-' * 50)
        
        # 添加内容
        content_para = doc.add_paragraph()
        content_para.add_run(content)
        
        # 添加页脚
        doc.add_page_break()
        footer = doc.add_paragraph('文档结束')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文档
        doc.save(filename)
        print(f"✓ 已创建: {filename}")
        return True
        
    except Exception as e:
        print(f"✗ 创建 {filename} 失败: {e}")
        return False

def main():
    """主函数"""
    print("开始生成三个Word文档...")
    
    # 文档定义
    documents = [
        {
            "title": "2026最新大模型发布",
            "content": """2026年，人工智能领域迎来了大模型技术的新一轮突破。各大科技巨头相继发布了新一代AI模型，这些模型在多个维度上实现了显著提升。

参数规模方面，新一代大模型普遍达到了万亿级别，但更重要的是架构优化带来的效率提升。新型的稀疏激活技术和混合专家模型（MoE）架构，使得模型在保持强大性能的同时，大幅降低了计算成本和能耗。

推理能力方面，2026年的模型展现出更强的逻辑推理和复杂问题解决能力。通过改进的注意力机制和训练方法，模型能够更好地理解上下文，进行多步推理。

多模态处理成为标配，新一代模型能够无缝处理文本、图像、音频和视频内容，实现真正的跨模态理解和生成。安全性和可控性也得到加强，内置了更完善的对齐机制和内容过滤系统。

这些技术进步使得大模型在医疗、教育、科研、创意等领域的应用更加广泛和深入，推动了人工智能技术的普惠化发展。""",
            "filename": "2026最新大模型发布.docx"
        },
        {
            "title": "大模型微调技术原理",
            "content": """大模型微调技术是将通用预训练模型适配到特定任务的关键方法。其核心原理是利用预训练模型已经学习到的丰富语言表示，通过少量领域数据调整模型参数，实现任务 specialization。

全参数微调是最直接的方法，更新模型所有权重参数。这种方法效果最佳，能够充分适应目标任务，但需要大量计算资源和训练数据，且容易导致灾难性遗忘。

部分参数微调是更高效的方案，如LoRA（低秩适应）技术。LoRA通过引入低秩分解矩阵来近似全参数更新，大幅减少了可训练参数数量（通常只有原模型的0.1%-1%），同时保持了接近全参数微调的性能。

适配器微调在模型中插入小型神经网络模块，只训练这些适配器，保持原始模型参数冻结。这种方法支持多任务学习，可以快速切换不同任务，计算效率极高。

提示微调（Prompt Tuning）通过优化输入提示的嵌入表示来调整模型行为，几乎不增加推理开销。前缀微调（Prefix Tuning）在每一层添加可训练的前缀向量，效果优于提示微调。

微调策略的选择需要综合考虑任务复杂度、数据量、计算资源和部署要求。在实际应用中，通常采用混合策略，结合多种微调方法的优势。""",
            "filename": "大模型微调技术原理.docx"
        },
        {
            "title": "怎么追求女孩",
            "content": """追求女孩是一个需要真诚、耐心和尊重的过程。成功的关键在于建立 genuine connection，而不是使用技巧或套路。

首先，展现真实的自己非常重要。不要为了迎合对方而改变本性，长期关系建立在 authenticity 基础上。自信但不自负，保持积极的生活态度和 personal growth。

沟通艺术至关重要。学会 active listening，真正关心对方的想法和感受。分享自己的经历和观点，但避免 dominating conversation。找到共同兴趣，创造有意义的对话。

尊重个人空间和边界。每个人都需要独处时间，过度热情可能造成压力。通过日常的小关心展现体贴，如记住重要日期、关注对方喜好，但避免 intrusive behavior。

时机把握需要 sensitivity。在适当的时候表达好感，但不要急于求成。观察对方的反应，尊重对方的选择。如果对方需要时间，给予足够的 space and patience。

处理 rejection 需要 maturity。如果对方没有相同感觉， gracefully accept 并保持尊重。真正的尊重体现在接受对方决定而不施加压力。无论结果如何，保持 dignity and kindness。

最重要的是，追求过程应该是 mutual and respectful。健康的关系建立在平等、理解和共同成长的基础上。""",
            "filename": "怎么追求女孩.docx"
        }
    ]
    
    # 创建文档
    success_count = 0
    for doc_info in documents:
        print(f"\n正在创建: {doc_info['title']}")
        if create_document(doc_info['title'], doc_info['content'], doc_info['filename']):
            success_count += 1
    
    # 输出结果
    print(f"\n{'='*50}")
    print(f"文档生成完成!")
    print(f"成功创建: {success_count}/{len(documents)} 个文档")
    print(f"\n生成的文件:")
    for doc_info in documents:
        file_exists = os.path.exists(doc_info['filename'])
        status = "✓" if file_exists else "✗"
        print(f"  {status} {doc_info['filename']}")
    
    if success_count == len(documents):
        print("\n所有文档已成功创建在workspace目录中!")
    else:
        print("\n部分文档创建失败，请检查错误信息。")

if __name__ == "__main__":
    main()